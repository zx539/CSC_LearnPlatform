from concurrent.futures import ThreadPoolExecutor
from datetime import datetime
from io import BytesIO
import json
import os
from pathlib import Path
import re
from threading import Lock
import time
from uuid import uuid4

from flask import Flask, jsonify, redirect, render_template, request, send_file, session, url_for
from requests import RequestException

from backend.services.multi_agent import MultiAgentLearningSystem
from backend.services.spark_client import SparkClient
from backend.services.user_store import UserStore


ROOT_DIR = Path(__file__).resolve().parents[1]
TEMPLATE_DIR = ROOT_DIR / "frontend" / "templates"
STATIC_DIR = ROOT_DIR / "frontend" / "static"
MATERIALS_DIR = ROOT_DIR / "materials"
ULTRA_CONFIG_PATH = MATERIALS_DIR / "星火SparkUltra-APIkey.txt"
LITE_CONFIG_PATH = MATERIALS_DIR / "星火SparkLite-APIkey.txt"
DOUBAO_CONFIG_PATH = MATERIALS_DIR / "豆包Seed2.0Pro-APIkey.txt"
DOUBAO_LITE_CONFIG_PATH = MATERIALS_DIR / "豆包Seed2.0Lite-APIkey.txt"
KB_DIR = ROOT_DIR / "data" / "knowledge_base"
OUTPUT_DIR = ROOT_DIR / "outputs"
LEGACY_USERS_DIR = ROOT_DIR / "data" / "users"
DEFAULT_USER_DATA_DIR = ROOT_DIR.parent / "csc_learnplatform_user_data"
USER_DATA_DIR = Path(os.getenv("USER_DATA_DIR", "")).expanduser() if os.getenv("USER_DATA_DIR") else DEFAULT_USER_DATA_DIR

app = Flask(__name__, template_folder=str(TEMPLATE_DIR), static_folder=str(STATIC_DIR))
app.secret_key = os.getenv("APP_SECRET_KEY", "software-cup-spark-secret")
user_store = UserStore(str(USER_DATA_DIR), legacy_dir=str(LEGACY_USERS_DIR))
generate_executor = ThreadPoolExecutor(max_workers=max(1, int(os.getenv("GENERATE_TASK_WORKERS", "2"))))
generate_tasks_lock = Lock()
generate_tasks: dict[str, dict] = {}
generate_task_ttl_seconds = max(300, int(os.getenv("GENERATE_TASK_TTL_SECONDS", "86400")))
generate_task_dir = USER_DATA_DIR / "_async_tasks"
generate_task_dir.mkdir(parents=True, exist_ok=True)


def _utc_now_iso() -> str:
    return datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ")


def _cleanup_generate_tasks():
    now_ts = time.time()
    with generate_tasks_lock:
        expired_ids = [
            task_id
            for task_id, task in generate_tasks.items()
            if task.get("status") in {"succeeded", "failed"}
            and now_ts - float(task.get("updated_ts", now_ts)) > generate_task_ttl_seconds
        ]
        for task_id in expired_ids:
            generate_tasks.pop(task_id, None)
    for task_file in generate_task_dir.glob("*.json"):
        try:
            payload = json.loads(task_file.read_text(encoding="utf-8"))
        except (OSError, ValueError):
            continue
        status = str(payload.get("status", "")).strip().lower()
        updated_ts = float(payload.get("updated_ts", now_ts))
        if status in {"succeeded", "failed"} and now_ts - updated_ts > generate_task_ttl_seconds:
            try:
                task_file.unlink(missing_ok=True)
            except OSError:
                continue


def _generate_task_path(task_id: str) -> Path:
    safe_task_id = re.sub(r"[^a-zA-Z0-9_-]", "", str(task_id or ""))
    return generate_task_dir / f"{safe_task_id}.json"


def _save_generate_task_to_file(task: dict):
    task_id = str(task.get("task_id", "")).strip()
    if not task_id:
        return
    path = _generate_task_path(task_id)
    tmp_path = path.with_suffix(".tmp")
    text = json.dumps(task, ensure_ascii=False, indent=2)
    tmp_path.write_text(text, encoding="utf-8")
    tmp_path.replace(path)


def _load_generate_task_from_file(task_id: str) -> dict | None:
    path = _generate_task_path(task_id)
    if not path.exists():
        return None
    try:
        payload = json.loads(path.read_text(encoding="utf-8"))
    except (OSError, ValueError):
        return None
    if not isinstance(payload, dict):
        return None
    return payload


def _set_generate_task(task_id: str, **updates) -> dict | None:
    now_ts = time.time()
    with generate_tasks_lock:
        task = generate_tasks.get(task_id)
        if not isinstance(task, dict):
            task = _load_generate_task_from_file(task_id)
            if isinstance(task, dict):
                generate_tasks[task_id] = task
        if not isinstance(task, dict):
            return None
        task.update(updates)
        task["updated_at"] = _utc_now_iso()
        task["updated_ts"] = now_ts
        snapshot = dict(task)
    _save_generate_task_to_file(snapshot)
    return snapshot


def _public_generate_task_payload(task: dict) -> dict:
    payload = {
        "task_id": task.get("task_id", ""),
        "status": task.get("status", "unknown"),
        "message": task.get("message", ""),
        "created_at": task.get("created_at", ""),
        "updated_at": task.get("updated_at", ""),
    }
    if task.get("status") == "succeeded":
        payload["result"] = task.get("result", {})
    elif task.get("status") == "failed":
        payload["error"] = task.get("error", "任务失败")
    return payload


def _run_generate_task(
    task_id: str,
    username: str,
    payload: dict,
    course: str,
    topic: str,
    dialogue: str,
    progress: str,
    model: str,
    run_name: str,
):
    _set_generate_task(task_id, status="running", message="任务已开始，正在生成学习画像与学习资源...")
    try:
        system = create_system(model=model)
        report = system.run(dialogue=dialogue, course=course, topic=topic, progress=progress)
        output_path, report_markdown = system.save_report(report, run_name)
        user_store.save_run(
            username=username,
            run_name=run_name,
            request_payload=payload,
            report=report,
            output_dir=str(output_path),
            report_markdown=report_markdown,
        )
        _set_generate_task(
            task_id,
            status="succeeded",
            message="生成完成",
            result={
                "run_name": run_name,
                "output_dir": str(output_path),
                "report": report,
                "report_markdown": report_markdown,
            },
        )
    except (FileNotFoundError, ValueError, RuntimeError, RequestException) as exc:
        _set_generate_task(task_id, status="failed", message="生成失败", error=str(exc))
    except Exception as exc:
        _set_generate_task(task_id, status="failed", message="生成失败", error=f"{type(exc).__name__}: {exc}")


def normalize_model(model: str) -> str:
    normalized = str(model or "").strip().lower()
    if normalized in {"lite", "sparklite", "4.0lite"}:
        return "lite"
    if normalized in {"doubao-lite", "doubao-seed-2-0-lite-260428", "seed2.0lite", "seed2-lite"}:
        return "doubao-seed-2-0-lite-260428"
    if normalized in {"doubao", "seed2.0pro", "seed2", "doubao-seed-2-0-pro-260215"}:
        return "doubao-seed-2-0-pro-260215"
    return "4.0Ultra"


def resolve_model_config_path(model: str) -> Path:
    normalized = normalize_model(model)
    if normalized == "lite":
        return LITE_CONFIG_PATH
    if normalized == "doubao-seed-2-0-lite-260428":
        return DOUBAO_LITE_CONFIG_PATH
    if normalized == "doubao-seed-2-0-pro-260215":
        return DOUBAO_CONFIG_PATH
    return ULTRA_CONFIG_PATH


def create_system(model: str = "4.0Ultra") -> MultiAgentLearningSystem:
    normalized_model = normalize_model(model)
    config_path = resolve_model_config_path(normalized_model)
    client = SparkClient(config_path=str(config_path), model=normalized_model)
    return MultiAgentLearningSystem(spark=client, kb_dir=str(KB_DIR), output_dir=str(OUTPUT_DIR))


def get_login_user() -> str | None:
    username = session.get("username")
    if isinstance(username, str) and username:
        return username
    return None


def require_login_json():
    user = get_login_user()
    if not user:
        return None, (jsonify({"error": "请先登录"}), 401)
    return user, None


def extract_profile_core(profile_payload: dict) -> dict:
    profile = profile_payload if isinstance(profile_payload, dict) else {}
    inner = profile.get("profile")
    if isinstance(inner, dict):
        nested_inner = inner.get("profile")
        if isinstance(nested_inner, dict):
            return nested_inner
        return inner
    return profile


def normalize_form_template(template: dict) -> dict:
    if not isinstance(template, dict):
        return {}
    questions = template.get("questions", [])
    if not isinstance(questions, list):
        questions = []
    normalized_questions: list[dict] = []
    for idx, item in enumerate(questions, start=1):
        if not isinstance(item, dict):
            continue
        qid = str(item.get("id", f"q{idx}")).strip() or f"q{idx}"
        question = str(item.get("question", "")).strip() or f"问题{idx}"
        raw_type = str(item.get("type", "")).strip().lower()
        options = item.get("options", [])
        if not isinstance(options, list):
            options = []
        options = [str(opt).strip() for opt in options if str(opt).strip()]
        q_type = "text"
        if raw_type in {"single", "single_choice", "radio", "单选"}:
            q_type = "single_choice"
        elif raw_type in {"multiple", "multi_choice", "checkbox", "多选"}:
            q_type = "multi_choice"
        elif raw_type in {"scale", "rating", "量表"}:
            q_type = "scale"
        elif options:
            if all(re.fullmatch(r"\d+(?:\.\d+)?", opt) for opt in options) and len(options) >= 3:
                q_type = "scale"
            elif any(keyword in question for keyword in ("哪些", "多选", "可多选")) or len(options) > 2:
                q_type = "multi_choice"
            else:
                q_type = "single_choice"
        normalized_questions.append(
            {
                "id": qid,
                "question": question,
                "type": q_type,
                "options": options,
                "required": bool(item.get("required", True)),
                "dimension": str(item.get("dimension", "学习进度")).strip() or "学习进度",
            }
        )
    normalized = dict(template)
    normalized["questions"] = normalized_questions
    return normalized


def normalize_report_forms(report: dict) -> dict:
    if not isinstance(report, dict):
        return {}
    report["progress_form_template"] = normalize_form_template(report.get("progress_form_template", {}))
    report["test_form_template"] = normalize_form_template(report.get("test_form_template", {}))
    return report


def _safe_str(value: object) -> str:
    return str(value or "").strip()


def _object_to_markdown(value: object, level: int = 0) -> str:
    if value is None:
        return "暂无"
    if isinstance(value, (str, int, float, bool)):
        return str(value)
    if isinstance(value, list):
        if not value:
            return "暂无"
        lines: list[str] = []
        for item in value:
            if isinstance(item, (dict, list)):
                lines.append(f"-\n{_object_to_markdown(item, level + 1)}")
            else:
                lines.append(f"- {item}")
        return "\n".join(lines)
    if isinstance(value, dict):
        if not value:
            return "暂无"
        indent = "  " * max(0, level)
        lines: list[str] = []
        for key, item in value.items():
            if isinstance(item, (dict, list)):
                lines.append(f"{indent}- **{key}**:\n{_object_to_markdown(item, level + 1)}")
            else:
                lines.append(f"{indent}- **{key}**: {item}")
        return "\n".join(lines)
    return str(value)


def _fallback_report_markdown(report_payload: dict) -> dict[str, str]:
    report = report_payload if isinstance(report_payload, dict) else {}
    resources = report.get("resources", {})
    if not isinstance(resources, dict):
        resources = {}
    resources_md = []
    for name, content in resources.items():
        text = content if isinstance(content, str) else _object_to_markdown(content)
        resources_md.append(f"### {name}\n\n{text}")
    return {
        "profile_md": f"## 学习画像\n\n{_object_to_markdown(report.get('profile', {}))}",
        "learning_path_md": f"## 学习路径\n\n{_object_to_markdown(report.get('learning_path', {}))}",
        "evaluation_md": f"## 学习评估\n\n{_object_to_markdown(report.get('evaluation', {'summary': '暂无学习评估。'}))}",
        "resources_md": "## 学习资源\n\n" + ("\n\n".join(resources_md) if resources_md else "暂无学习资源。"),
    }


def _compose_report_markdown(report_payload: dict, report_markdown_payload: dict) -> str:
    report_markdown = report_markdown_payload if isinstance(report_markdown_payload, dict) else {}
    fallback = _fallback_report_markdown(report_payload)
    resources_md = _safe_str(report_markdown.get("resources_md", ""))
    if not resources_md:
        resources = report_payload.get("resources", {}) if isinstance(report_payload, dict) else {}
        if isinstance(resources, dict) and resources:
            resource_sections = []
            for name, content in resources.items():
                text = content if isinstance(content, str) else _object_to_markdown(content)
                resource_sections.append(f"### {name}\n\n{text}")
            resources_md = "## 学习资源\n\n" + "\n\n".join(resource_sections)
    return "\n\n".join(
        [
            _safe_str(report_markdown.get("profile_md", "")) or fallback["profile_md"],
            resources_md or fallback["resources_md"],
            _safe_str(report_markdown.get("learning_path_md", "")) or fallback["learning_path_md"],
            _safe_str(report_markdown.get("evaluation_md", "")) or fallback["evaluation_md"],
        ]
    ).strip()


def _markdown_to_docx_bytes(markdown_text: str) -> bytes:
    from docx import Document

    doc = Document()
    in_code_block = False
    for raw_line in str(markdown_text or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
        if in_code_block:
            doc.add_paragraph(line, style="Intense Quote")
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(4, len(heading_match.group(1)))
            doc.add_heading(heading_match.group(2).strip(), level=level)
            continue
        if re.match(r"^\d+\.\s+", stripped):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)
    buff = BytesIO()
    doc.save(buff)
    return buff.getvalue()


def _markdown_to_pdf_bytes(markdown_text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import mm
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.platypus import Paragraph, Preformatted, SimpleDocTemplate, Spacer

    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    styles = getSampleStyleSheet()
    for style_name in ("BodyText", "Heading1", "Heading2", "Heading3", "Code"):
        if style_name in styles:
            styles[style_name].fontName = "STSong-Light"
    story = []
    in_code_block = False
    for raw_line in str(markdown_text or "").splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()
        if stripped.startswith("```"):
            in_code_block = not in_code_block
            story.append(Spacer(1, 4))
            continue
        if in_code_block:
            story.append(Preformatted(line, styles["Code"]))
            continue
        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = min(3, len(heading_match.group(1)))
            style_name = f"Heading{level}" if level > 0 else "Heading1"
            story.append(Paragraph(heading_match.group(2).strip(), styles[style_name]))
            story.append(Spacer(1, 6))
            continue
        if stripped.startswith("- ") or stripped.startswith("* "):
            story.append(Paragraph(f"• {stripped[2:].strip()}", styles["BodyText"]))
            continue
        if re.match(r"^\d+\.\s+", stripped):
            story.append(Paragraph(stripped, styles["BodyText"]))
            continue
        story.append(Paragraph(stripped if stripped else "&nbsp;", styles["BodyText"]))
    buff = BytesIO()
    doc = SimpleDocTemplate(buff, pagesize=A4, leftMargin=14 * mm, rightMargin=14 * mm, topMargin=14 * mm, bottomMargin=14 * mm)
    doc.build(story)
    return buff.getvalue()


@app.get("/login")
def login_page():
    if get_login_user():
        return redirect(url_for("index"))
    return render_template("login.html")


@app.get("/forgot-password")
def forgot_password_page():
    if get_login_user():
        return redirect(url_for("index"))
    return render_template("forgot_password.html")


@app.post("/login")
def login_action():
    username = str(request.form.get("username", "")).strip()
    password = str(request.form.get("password", "")).strip()
    if not username or not password:
        return render_template("login.html", error="用户名和密码不能为空")

    try:
        ok, created = user_store.login_or_register(username, password)
    except ValueError as exc:
        return render_template("login.html", error=str(exc))
    if not ok:
        return render_template("login.html", error="用户名或密码错误")

    session["username"] = username
    session["progress_submitted_this_login"] = False
    if created:
        return redirect(url_for("index", notice="首次登录已创建账号"))
    return redirect(url_for("index"))


@app.post("/api/auth/security-questions")
def auth_security_questions():
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    username = str(payload.get("username", "")).strip()
    if not username:
        return jsonify({"error": "用户名不能为空"}), 400
    try:
        questions = user_store.get_security_questions(username)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"username": username, "questions": questions})


@app.post("/api/auth/reset-password")
def auth_reset_password():
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    username = str(payload.get("username", "")).strip()
    answers_payload = payload.get("answers", [])
    new_password = str(payload.get("new_password", "")).strip()
    if not username:
        return jsonify({"error": "用户名不能为空"}), 400
    if not isinstance(answers_payload, list):
        return jsonify({"error": "answers 必须是数组类型"}), 400
    answers = [str(item).strip() for item in answers_payload]
    try:
        user_store.reset_password_by_security_questions(username, answers, new_password)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"message": "密码已重置，请使用新密码登录"})


@app.get("/logout")
def logout():
    session.clear()
    return redirect(url_for("login_page"))


@app.get("/")
def index():
    username = get_login_user()
    if not username:
        return redirect(url_for("login_page"))
    return render_template("index.html", username=username, notice=request.args.get("notice", ""))


@app.post("/api/generate")
def generate():
    username, err = require_login_json()
    if err:
        return err

    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400

    course = str(payload.get("course", "")).strip()
    topic = str(payload.get("topic", "")).strip()
    dialogue = str(payload.get("dialogue", "")).strip()
    progress = str(payload.get("progress", "")).strip()
    model = str(payload.get("model", "4.0Ultra")).strip() or "4.0Ultra"

    if not course or not topic or not dialogue:
        return jsonify({"error": "course、topic、dialogue 为必填项"}), 400

    run_name = datetime.now().strftime("run_%Y%m%d_%H%M%S")
    task_id = uuid4().hex
    now_iso = _utc_now_iso()
    now_ts = time.time()
    with generate_tasks_lock:
        generate_tasks[task_id] = {
            "task_id": task_id,
            "username": username,
            "status": "queued",
            "message": "任务已提交，等待执行",
            "created_at": now_iso,
            "updated_at": now_iso,
            "created_ts": now_ts,
            "updated_ts": now_ts,
            "result": None,
            "error": "",
        }
        task_snapshot = dict(generate_tasks[task_id])
    _save_generate_task_to_file(task_snapshot)
    generate_executor.submit(
        _run_generate_task,
        task_id,
        username,
        payload,
        course,
        topic,
        dialogue,
        progress,
        model,
        run_name,
    )
    _cleanup_generate_tasks()
    return jsonify({"task_id": task_id, "status": "queued", "message": "任务已提交，请轮询任务状态"})


@app.get("/api/generate/<task_id>")
def get_generate_task(task_id: str):
    username, err = require_login_json()
    if err:
        return err

    with generate_tasks_lock:
        task = generate_tasks.get(task_id)
        task_copy = dict(task) if isinstance(task, dict) else None
    if not isinstance(task_copy, dict):
        task_copy = _load_generate_task_from_file(task_id)
        if isinstance(task_copy, dict):
            with generate_tasks_lock:
                generate_tasks[task_id] = dict(task_copy)
    if not isinstance(task_copy, dict):
        return jsonify({"error": "任务不存在或已过期"}), 404
    if str(task_copy.get("username", "")) != username:
        return jsonify({"error": "无权访问该任务"}), 403
    _cleanup_generate_tasks()
    return jsonify(_public_generate_task_payload(task_copy))


@app.post("/api/tutor")
def tutor():
    username, err = require_login_json()
    if err:
        return err

    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400

    question = str(payload.get("question", "")).strip()
    topic = str(payload.get("topic", "")).strip()
    model = str(payload.get("model", "4.0Ultra")).strip() or "4.0Ultra"
    profile = payload.get("profile")
    memory_payload = payload.get("memory", [])
    if not isinstance(profile, dict):
        return jsonify({"error": "profile 必须是对象类型"}), 400
    if not isinstance(memory_payload, list):
        return jsonify({"error": "memory 必须是数组类型"}), 400
    memory: list[dict[str, str]] = []
    for item in memory_payload[:20]:
        if not isinstance(item, dict):
            continue
        q = str(item.get("question", "")).strip()
        a = str(item.get("answer", "")).strip()
        if not q and not a:
            continue
        memory.append({"question": q, "answer": a})
    if not question or not topic:
        return jsonify({"error": "question、topic 为必填项"}), 400

    try:
        system = create_system(model=model)
        answer = system.tutor(question=question, profile=profile, topic=topic, memory=memory)
        user_store.save_tutor(username=username, question=question, answer=answer, topic=topic)
    except (FileNotFoundError, ValueError, RuntimeError, RequestException) as exc:
        return jsonify({"error": str(exc)}), 500

    return jsonify({"answer": answer})


@app.get("/api/user/profile")
def user_profile():
    username, err = require_login_json()
    if err:
        return err

    profile = user_store.get_display_profile(username)
    return jsonify(
        {
            "username": profile["username"],
            "nickname": profile["nickname"],
            "runs": user_store.list_runs(username),
            "latest": user_store.get_latest_report(username),
            "progress_logs": user_store.get_progress_logs(username, limit=10),
        }
    )


@app.get("/api/user/center")
def user_center():
    username, err = require_login_json()
    if err:
        return err
    center = user_store.get_user_center(username)
    center["avatar_url"] = "/api/user/avatar"
    return jsonify(center)


@app.post("/api/user/nickname")
def user_nickname():
    username, err = require_login_json()
    if err:
        return err
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    nickname = str(payload.get("nickname", "")).strip()
    if not nickname:
        return jsonify({"error": "昵称不能为空"}), 400
    try:
        user_store.set_nickname(username, nickname)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"message": "昵称已更新", "nickname": nickname})


@app.get("/api/user/avatar")
def user_avatar():
    username, err = require_login_json()
    if err:
        return err
    avatar_path = user_store.get_avatar(username)
    if not avatar_path:
        return jsonify({"error": "未设置头像"}), 404
    suffix = avatar_path.suffix.lower()
    mimetype = "image/png"
    if suffix in {".jpg", ".jpeg"}:
        mimetype = "image/jpeg"
    elif suffix == ".webp":
        mimetype = "image/webp"
    return send_file(avatar_path, mimetype=mimetype)


@app.post("/api/user/avatar")
def upload_user_avatar():
    username, err = require_login_json()
    if err:
        return err
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    avatar_data = str(payload.get("avatar_data", "")).strip()
    if not avatar_data:
        return jsonify({"error": "头像数据不能为空"}), 400
    try:
        user_store.save_avatar(username, avatar_data)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"message": "头像已更新", "avatar_url": "/api/user/avatar"})


@app.post("/api/user/change-password")
def user_change_password():
    username, err = require_login_json()
    if err:
        return err
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    current_password = str(payload.get("current_password", "")).strip()
    new_password = str(payload.get("new_password", "")).strip()
    if not current_password or not new_password:
        return jsonify({"error": "当前密码和新密码不能为空"}), 400
    try:
        user_store.change_password(username, current_password, new_password)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"message": "密码修改成功"})


@app.post("/api/user/security-questions")
def user_security_questions():
    username, err = require_login_json()
    if err:
        return err
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    current_password = str(payload.get("current_password", "")).strip()
    questions = payload.get("questions", [])
    if not current_password:
        return jsonify({"error": "当前密码不能为空"}), 400
    if not isinstance(questions, list):
        return jsonify({"error": "questions 必须是数组类型"}), 400
    try:
        user_store.set_security_questions(username, current_password, questions)
    except ValueError as exc:
        return jsonify({"error": str(exc)}), 400
    return jsonify({"message": "密保问题已保存"})


@app.get("/api/projects")
def list_projects():
    username, err = require_login_json()
    if err:
        return err

    return jsonify({"projects": user_store.list_runs(username)})


@app.get("/api/user/run/<run_name>")
def user_run(run_name: str):
    username, err = require_login_json()
    if err:
        return err

    payload = user_store.get_run(username, run_name)
    if payload is None:
        return jsonify({"error": "未找到该历史记录"}), 404
    report = payload.get("report")
    if isinstance(report, dict):
        payload["report"] = normalize_report_forms(report)
    return jsonify(payload)


@app.get("/api/user/run/<run_name>/document")
def download_run_document(run_name: str):
    username, err = require_login_json()
    if err:
        return err
    fmt = str(request.args.get("format", "pdf")).strip().lower()
    if fmt not in {"pdf", "docx"}:
        return jsonify({"error": "仅支持 pdf 或 docx"}), 400
    payload = user_store.get_run(username, run_name)
    if payload is None:
        return jsonify({"error": "未找到该历史记录"}), 404
    report = payload.get("report", {})
    report_markdown = payload.get("report_markdown", {})
    merged_markdown = _compose_report_markdown(report, report_markdown)
    if not merged_markdown:
        return jsonify({"error": "当前项目暂无可导出的学习内容"}), 400
    now_tag = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_run_name = re.sub(r"[^0-9A-Za-z_\-\u4e00-\u9fa5]", "_", run_name).strip("_") or "learning_project"
    try:
        if fmt == "docx":
            data = _markdown_to_docx_bytes(merged_markdown)
            return send_file(
                BytesIO(data),
                mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                as_attachment=True,
                download_name=f"{safe_run_name}_{now_tag}.docx",
            )
        data = _markdown_to_pdf_bytes(merged_markdown)
        return send_file(
            BytesIO(data),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=f"{safe_run_name}_{now_tag}.pdf",
        )
    except ModuleNotFoundError:
        return jsonify({"error": "导出依赖未安装，请先安装 requirements.txt"}), 500


@app.delete("/api/user/run/<run_name>")
def delete_user_run(run_name: str):
    username, err = require_login_json()
    if err:
        return err

    ok = user_store.delete_run(username, run_name)
    if not ok:
        return jsonify({"error": "未找到该历史记录"}), 404

    return jsonify({"message": "历史记录已删除", "run_name": run_name})


@app.get("/api/share/export/<run_name>")
def export_run_share_file(run_name: str):
    username, err = require_login_json()
    if err:
        return err
    payload = user_store.get_run(username, run_name)
    if payload is None:
        return jsonify({"error": "未找到该历史记录"}), 404
    profile = user_store.get_display_profile(username)
    share_payload = {
        "format": "csc_learnplatform_share_v1",
        "shared_at": _utc_now_iso(),
        "shared_by": profile,
        "project": {
            "source_run_name": run_name,
            "run_payload": payload,
        },
    }
    data = json.dumps(share_payload, ensure_ascii=False, indent=2).encode("utf-8")
    safe_run_name = re.sub(r"[^0-9A-Za-z_\-\u4e00-\u9fa5]", "_", run_name).strip("_") or "learning_project"
    return send_file(
        BytesIO(data),
        mimetype="application/json",
        as_attachment=True,
        download_name=f"{safe_run_name}_share.json",
    )


@app.post("/api/share/import")
def import_run_share_file():
    username, err = require_login_json()
    if err:
        return err
    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    share_payload = payload.get("share_payload")
    if not isinstance(share_payload, dict):
        return jsonify({"error": "share_payload 必须是对象类型"}), 400
    if str(share_payload.get("format", "")).strip() != "csc_learnplatform_share_v1":
        return jsonify({"error": "分享文件格式不正确"}), 400
    project_payload = share_payload.get("project", {})
    if not isinstance(project_payload, dict):
        return jsonify({"error": "分享文件缺少 project 信息"}), 400
    run_payload = project_payload.get("run_payload", {})
    if not isinstance(run_payload, dict):
        return jsonify({"error": "分享文件缺少 run_payload"}), 400
    request_payload = run_payload.get("request", {})
    report = run_payload.get("report", {})
    report_markdown = run_payload.get("report_markdown", {})
    if not isinstance(request_payload, dict) or not isinstance(report, dict):
        return jsonify({"error": "分享文件中的项目数据不完整"}), 400
    if not isinstance(report_markdown, dict):
        report_markdown = {}
    source_run_name = str(project_payload.get("source_run_name", "")).strip() or str(run_payload.get("run_name", "")).strip()
    base_name = datetime.now().strftime("run_share_%Y%m%d_%H%M%S")
    run_name = base_name
    suffix = 1
    while user_store.get_run(username, run_name) is not None:
        suffix += 1
        run_name = f"{base_name}_{suffix}"
    shared_by_payload = share_payload.get("shared_by", {})
    shared_by = {
        "username": str(shared_by_payload.get("username", "")).strip(),
        "nickname": str(shared_by_payload.get("nickname", "")).strip(),
    }
    user_store.save_run(
        username=username,
        run_name=run_name,
        request_payload=request_payload,
        report=report,
        output_dir="",
        report_markdown=report_markdown,
        shared_by=shared_by,
        source_run_name=source_run_name,
    )
    return jsonify(
        {
            "message": "分享项目已导入",
            "run_name": run_name,
            "shared_by": shared_by,
            "source_run_name": source_run_name,
        }
    )


@app.post("/api/progress/checkin")
def progress_checkin():
    username, err = require_login_json()
    if err:
        return err

    payload = request.get_json(silent=True)
    if payload is None:
        return jsonify({"error": "请求体必须是JSON"}), 400
    checkin = payload.get("checkin")
    if not isinstance(checkin, dict):
        return jsonify({"error": "checkin 必须是对象类型"}), 400
    form_type = str(payload.get("form_type", "progress")).strip().lower() or "progress"
    if form_type not in {"progress", "test"}:
        return jsonify({"error": "form_type 仅支持 progress 或 test"}), 400
    responses = checkin.get("responses")
    if not isinstance(responses, list) or not responses:
        return jsonify({"error": "问卷答案不能为空，请先完成问卷填写"}), 400

    run_name = str(payload.get("run_name", "")).strip()
    latest = user_store.get_latest_report(username)
    target_payload = None
    if run_name:
        target_payload = user_store.get_run(username, run_name)
    elif isinstance(latest, dict):
        target_payload = latest

    if not isinstance(target_payload, dict):
        return jsonify({"error": "未找到对应学习项目，请先选择已有项目"}), 400

    report = target_payload.get("report")
    if not isinstance(report, dict):
        return jsonify({"error": "历史报告格式错误"}), 500
    report = normalize_report_forms(report)
    target_payload["report"] = report

    profile = report.get("profile")
    path = report.get("learning_path")
    if not isinstance(profile, dict) or not isinstance(path, dict):
        return jsonify({"error": "历史报告缺少学习画像或学习路径"}), 500
    profile_core = extract_profile_core(profile)

    req_payload = target_payload.get("request", {})
    default_model = "4.0Ultra"
    if isinstance(req_payload, dict):
        default_model = str(req_payload.get("model", "4.0Ultra")).strip() or "4.0Ultra"
    model = str(payload.get("model", default_model)).strip() or default_model

    system = create_system(model=model)
    try:
        evaluation = system.evaluate_learning(progress_payload=checkin, profile=profile_core, path=path)
    except (FileNotFoundError, ValueError, RuntimeError, RequestException) as exc:
        return jsonify({"error": str(exc)}), 500
    report["evaluation"] = evaluation
    generated_next_form = False

    questionnaire_history = report.get("questionnaire_history", [])
    if not isinstance(questionnaire_history, list):
        questionnaire_history = []

    if form_type == "test":
        req_topic = ""
        req_data = target_payload.get("request", {})
        if isinstance(req_data, dict):
            req_topic = str(req_data.get("topic", "")).strip()
        topic = req_topic or str(checkin.get("topic", "")).strip()
        current_progress_form = report.get("progress_form_template", {})
        if not isinstance(current_progress_form, dict):
            current_progress_form = {}
        current_test_form = report.get("test_form_template", {})
        if not isinstance(current_test_form, dict):
            current_test_form = {}
        stage_state = report.get("stage_state", {})
        stage_no = 1
        if isinstance(stage_state, dict):
            try:
                stage_no = max(1, int(stage_state.get("current_stage_no", 1)))
            except (TypeError, ValueError):
                stage_no = 1
        try:
            stage_no = max(1, int(current_progress_form.get("stage_no", stage_no)))
        except (TypeError, ValueError):
            stage_no = max(1, stage_no)
        try:
            stage_no = max(1, int(current_test_form.get("stage_no", stage_no)))
        except (TypeError, ValueError):
            stage_no = max(1, stage_no)
        try:
            next_progress_form = system.build_progress_form(
                topic=topic,
                path=path,
                profile=profile_core,
                stage_no=stage_no,
                last_checkin=checkin,
                last_evaluation=evaluation,
            )
            next_test_form = system.build_test_form(
                topic=topic,
                path=path,
                profile=profile_core,
                stage_no=stage_no,
                last_checkin=checkin,
                last_evaluation=evaluation,
            )
        except (FileNotFoundError, ValueError, RuntimeError, RequestException) as exc:
            return jsonify({"error": str(exc)}), 500
        report["progress_form_template"] = next_progress_form
        report["test_form_template"] = next_test_form
        generated_next_form = True
        if current_progress_form:
            questionnaire_history.insert(
                0,
                {
                    "form_type": "progress",
                    "stage_no": stage_no,
                    "created_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
                    "version_tag": "old",
                    "form": current_progress_form,
                    "markdown": system._progress_form_to_markdown(current_progress_form),
                },
            )
        questionnaire_history.insert(
            0,
            {
                "form_type": "progress",
                "stage_no": stage_no,
                "created_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
                "version_tag": "new",
                "form": next_progress_form,
                "markdown": system._progress_form_to_markdown(next_progress_form),
            },
        )
        if current_test_form:
            questionnaire_history.insert(
                0,
                {
                    "form_type": "test",
                    "stage_no": stage_no,
                    "created_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
                    "version_tag": "old",
                    "form": current_test_form,
                    "markdown": system._progress_form_to_markdown(current_test_form),
                },
            )
        questionnaire_history.insert(
            0,
            {
                "form_type": "test",
                "stage_no": stage_no,
                "created_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
                "version_tag": "new",
                "form": next_test_form,
                "markdown": system._progress_form_to_markdown(next_test_form),
            },
        )

    report["questionnaire_history"] = questionnaire_history[:200]
    report_markdown = system.build_report_markdown(report)
    target_payload["report"] = report
    target_payload["report_markdown"] = report_markdown
    output_dir = str(target_payload.get("output_dir", "")).strip()
    if output_dir:
        system.persist_markdown_files(Path(output_dir), report_markdown)
        system.persist_questionnaire_markdown_files(Path(output_dir), report.get("questionnaire_history", []))

    progress_entry = {
        "created_at": datetime.utcnow().strftime("%Y-%m-%dT%H:%M:%SZ"),
        "run_name": target_payload.get("run_name", ""),
        "form_type": form_type,
        "checkin": checkin,
        "evaluation": evaluation,
    }
    latest_checkins = target_payload.get("progress_checkins", [])
    if not isinstance(latest_checkins, list):
        latest_checkins = []
    latest_checkins.insert(0, progress_entry)
    target_payload["progress_checkins"] = latest_checkins[:100]

    target_run_name = str(target_payload.get("run_name", "")).strip()
    if target_run_name:
        user_store.update_run(username, target_run_name, target_payload)
    if isinstance(latest, dict) and str(latest.get("run_name", "")).strip() == target_run_name:
        user_store.update_latest_report(username, target_payload)
    user_store.append_progress_log(username, progress_entry)
    session["progress_submitted_this_login"] = True

    return jsonify(
        {
            "message": "问卷已提交",
            "evaluation": evaluation,
            "evaluation_md": report_markdown.get("evaluation_md", ""),
            "report_markdown": report_markdown,
            "run_name": target_run_name,
            "next_progress_form": report.get("progress_form_template", {}),
            "next_test_form": report.get("test_form_template", {}),
            "generated_next_form": generated_next_form,
            "form_type": form_type,
        }
    )


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=8000, debug=True)
